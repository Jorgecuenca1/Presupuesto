"""Genera el Manual de Usuario en formato Word (.docx) para el aplicativo de Presupuesto."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x59, 0x59, 0x59)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def h1(doc, txt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = AZUL


def h2(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = AZUL


def h3(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = GRIS


def parrafo(doc, txt, negrita=False):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.size = Pt(11)
    if negrita:
        r.bold = True
    return p


def lista(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(it).font.size = Pt(11)


def tabla_func(doc, filas):
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    hdr[0].text = 'Función'
    hdr[1].text = 'Descripción'
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_bg(c, '1F4E79')
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for f, d in filas:
        row = t.add_row().cells
        row[0].text = f
        row[1].text = d
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


# ===== Documento =====
doc = Document()

# Configurar márgenes
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Estilo base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Portada
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\n\n\nMANUAL DE USUARIO')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = AZUL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Aplicativo de Presupuesto Municipal')
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = GRIS

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Alcaldía de Puerto López — Meta, Colombia')
r.font.size = Pt(14)
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Secretaría de Hacienda')
r.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Vigencia 2026')
r.bold = True
r.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Versión 1.0')
r.font.size = Pt(11)

doc.add_page_break()

# ===== 1. DESCRIPCIÓN DEL PROGRAMA =====
h1(doc, '1. Descripción del Programa')

parrafo(doc,
    'El Aplicativo de Presupuesto es un sistema web desarrollado en Django (Python) para la '
    'Alcaldía de Puerto López (Meta, Colombia). Permite elaborar, proyectar y administrar el '
    'presupuesto anual del municipio de acuerdo con la normatividad colombiana vigente: '
    'CUIPO (Clasificador Único de Ingresos y Gastos), Ley 617 de 2000, Ley 44 de 1990 '
    '(predial), Ley 14 de 1983 (ICA), Sistema General de Participaciones (SGP) y Plan '
    'Operativo Anual de Inversiones (POAI).'
)

parrafo(doc,
    'El programa centraliza en una sola herramienta la proyección de ingresos '
    '(predial, industria y comercio, avisos y tableros, estampillas, transferencias), '
    'la programación de gastos de funcionamiento, deuda pública e inversión, y la '
    'generación de los anexos oficiales del presupuesto municipal.'
)

h2(doc, '1.1 Objetivo')
parrafo(doc,
    'Automatizar los cálculos tributarios y presupuestales para producir los Anexos 1 '
    '(Ingresos) y 2 (Gastos) del presupuesto, minimizar errores manuales en hojas de '
    'cálculo y mantener la trazabilidad de los parámetros de proyección.'
)

h2(doc, '1.2 Alcance')
lista(doc, [
    'Proyección de impuesto predial urbano y rural (vigencia actual y vigencias anteriores).',
    'Proyección del impuesto de Industria y Comercio (ICA), Avisos y Tableros.',
    'Cálculo de transferencias SGP y rubros por métodos IPC, ICN y POAI.',
    'Programación de gastos: funcionamiento, servicio a la deuda e inversión.',
    'Control de límites Ley 617/2000 (Concejo, Personería, Administración Central).',
    'Cálculo de TCPA (Tasa de Crecimiento Promedio Anual) sobre cifras históricas CUIPO.',
    'Registro de vigencias futuras y costos de personal (planta).',
    'Exportación a Excel de reportes e importación masiva desde archivos XLSX.',
])

h2(doc, '1.3 Tecnología')
lista(doc, [
    'Backend: Django (Python 3).',
    'Base de datos: SQLite (db.sqlite3).',
    'Frontend: plantillas HTML con Bootstrap.',
    'Importación/exportación: openpyxl.',
    'Despliegue: servidor local con manage.py runserver o contenedor Docker.',
])

h2(doc, '1.4 Estructura de módulos')
tabla_func(doc, [
    ('core', 'Parámetros del sistema, autenticación (login/registro), dashboard, tablas Ley 617/2000.'),
    ('ingresos', 'Predial, ICA, Avisos y Tableros, rubros de ingreso, cartera, cifras históricas, TCPA.'),
    ('gastos', 'Secciones, rubros de gasto, ejecución, deuda pública, personal, vigencias futuras, techos de inversión.'),
    ('templates', 'Vistas HTML de todos los módulos.'),
    ('static', 'Recursos estáticos (CSS, JS).'),
])

doc.add_page_break()

# ===== 2. ACCESO Y PARÁMETROS INICIALES =====
h1(doc, '2. Acceso al Sistema')

h2(doc, '2.1 Iniciar el servidor')
parrafo(doc, 'Desde la carpeta del proyecto, abrir una consola y ejecutar:')
p = doc.add_paragraph()
r = p.add_run('python manage.py runserver')
r.font.name = 'Consolas'
r.font.size = Pt(11)
parrafo(doc, 'Luego abrir el navegador en: http://127.0.0.1:8000/')

h2(doc, '2.2 Ingreso')
lista(doc, [
    'URL de login: /login/',
    'Registro de usuarios nuevos: /registro/',
    'Panel principal (Dashboard): / (después de autenticarse).',
])

h2(doc, '2.3 Parámetros del Sistema (antes de calcular)')
parrafo(doc,
    'Antes de ejecutar cualquier cálculo debe configurar los parámetros de la vigencia '
    'desde el menú "Parámetros" (/parametros/):'
)
tabla_func(doc, [
    ('Vigencia Fiscal', 'Año para el que se proyecta el presupuesto (ej. 2026).'),
    ('Valor UVT ($)', 'Valor de la UVT publicado por la DIAN para la vigencia.'),
    ('Tasa IPC', 'Inflación proyectada (ej. 0.051 = 5.1%).'),
    ('Tasa Crecimiento ICN', 'Tasa de crecimiento de Ingresos Corrientes de la Nación.'),
    ('Tasa PIB Nominal', 'PIB nominal proyectado (se aplica a ICA).'),
    ('% Crecimiento Viviendas', 'Incremento anual de viviendas nuevas, aplicado a UV (Urbano Vivienda). Por defecto 1.5%.'),
    ('% Eficiencia de Recaudo', 'Porcentaje global de recaudo efectivo (antes llamado Cultura de Pago). Por defecto 70%.'),
    ('% Base Recaudo Cartera', 'Porcentaje base sobre el valor de cartera de vigencias anteriores (por defecto 40%).'),
    ('% Urbano / % Rural Cartera', 'Distribución de la cartera entre urbano (10%) y rural (90%).'),
    ('POAI Total Inversión', 'Monto total de inversión (sin Educación, Vivienda, Salud) para rubros con método POAI.'),
    ('Categoría del Municipio', 'Categoría según Ley 617/2000 (Puerto López = 6 / Sexta).'),
    ('Valor SMLMV', 'Salario mínimo mensual legal vigente para cálculos de Concejo y Personería.'),
])

doc.add_page_break()

# ===== 3. FUNCIONES DEL PROGRAMA =====
h1(doc, '3. Funciones del Programa')

# -------- 3.1 Ingresos
h2(doc, '3.1 Módulo de Ingresos')

h3(doc, '3.1.1 Tarifas Predial')
parrafo(doc,
    'Se registran las tarifas por mil (‰) por categoría y rango de UVT. Categorías disponibles:'
)
lista(doc, [
    'UV – Urbano Vivienda',
    'UEF – Urbano Edificado Actividades Financieras',
    'UED – Urbano Edificado Demás',
    'UNEU – Urbano No Edificado Urbanizable No Urbanizado',
    'UNUE – Urbano No Edificado Urbanizado No Edificado',
    'UNNU – Urbano No Edificado No Urbanizable',
    'RU – Rural',
    'PE – Parcelación/Finca Recreo Edificado',
    'PNE – Parcelación/Finca Recreo No Edificado',
])
parrafo(doc,
    'Nota: Los urbanos no edificados (UNEU/UNUE/UNNU) se agrupan con Rural en los reportes, '
    'conservando su tarifa propia.'
)

h3(doc, '3.1.2 Contribuyentes Predial e Importación')
lista(doc, [
    'Alta manual de predios: dirección, propietario, cédula catastral, avalúo y categoría.',
    'Importación masiva desde el archivo "Tabla Predial Comparativo con vigencia anterior.xlsx".',
    'Clasificación automática según columnas TIPO / DESTINO / CLASE del catastro.',
    'Los predios con USO PUBLICO, INFRAESTRUCTURA, FUNERARIOS o CULTURAL se reclasifican a RU (exentos urbanos).',
])

h3(doc, '3.1.3 Cálculo de Predial')
parrafo(doc, 'Fórmulas aplicadas:', negrita=True)
lista(doc, [
    'Recaudo Potencial = Σ(Avalúo × Tarifa/1000) por rango UVT y categoría.',
    'Proyección = Recaudo Potencial × % Eficiencia × (1 + % Crecimiento Viviendas)  →  el factor (1+%) aplica SOLO a categoría UV.',
    'Predial vigencias anteriores = Valor Cartera × (% Base/100) × (% Urbano o % Rural/100).',
])
parrafo(doc,
    'Pantalla "Cálculo Predial": muestra tabla por categoría con total de avalúo, tarifa, '
    'recaudo potencial, % eficiencia, proyección y cantidad de predios. Fila de totales por '
    'columna al final.'
)

h3(doc, '3.1.4 Tarifas e ICA (Industria y Comercio)')
parrafo(doc,
    'Tarifas por código de actividad (101 Industrial, 201-204 Comercial, 301-302 Servicios, '
    '401 Financiero).'
)
parrafo(doc, 'Fórmula ICA:', negrita=True)
lista(doc, [
    'Ingresos Proyectados = Ingresos Brutos × (1 + PIB Nominal).',
    'Impuesto = Ingresos Proyectados × Tarifa/1000.',
    'Avisos y Tableros = 15% del total ICA.',
])

h3(doc, '3.1.5 Rubros de Ingreso y Métodos de Cálculo')
tabla_func(doc, [
    ('PUVA', 'Predial Urbano Vigencia Actual (suma la proyección urbana).'),
    ('PUAN', 'Predial Urbano Vigencias Anteriores (cartera).'),
    ('PRVA', 'Predial Rural Vigencia Actual.'),
    ('PRAN', 'Predial Rural Vigencias Anteriores.'),
    ('ICAI / ICAC / ICAS', 'ICA Industrial / Comercial / Servicios.'),
    ('AT', 'Avisos y Tableros (15% del ICA).'),
    ('IPC', 'Recaudo anterior × (1 + IPC).'),
    ('ICN', 'Recaudo anterior × (1 + tasa ICN).'),
    ('POAI', 'POAI Inversión × tarifa del rubro.'),
    ('MAN', 'Valor ingresado manualmente.'),
])

h3(doc, '3.1.6 Cálculo Integral')
parrafo(doc,
    'El botón "Calcular Todos" recorre todos los rubros, ejecuta predial, ICA y asigna '
    'los valores según el método de cada rubro. Los títulos/subtotales se recalculan '
    'automáticamente de abajo hacia arriba por nivel jerárquico.'
)

h3(doc, '3.1.7 Cifras Históricas y TCPA')
lista(doc, [
    'Permite importar cifras históricas 2022-2025 desde reportes CUIPO.',
    'Calcula TCPA (Tasa Compuesta Promedio Anual) de ingresos y gastos.',
    'Marca ICLD, SGP y SGP Libre Asignación para los reportes de Ley 617.',
])

h3(doc, '3.1.8 Reporte y Exportación')
lista(doc, [
    'Reporte jerárquico de ingresos por código, descripción, fuente y apropiación.',
    'Exportación a Excel conservando la estructura del Anexo 1.',
])

# -------- 3.2 Gastos
h2(doc, '3.2 Módulo de Gastos')

h3(doc, '3.2.1 Secciones y Fuentes')
lista(doc, [
    'Alta de secciones presupuestales (Central, Concejo, Personería, etc.).',
    'Catálogo de fuentes de financiación (SGP, Recursos Propios, Estampillas, etc.).',
])

h3(doc, '3.2.2 Rubros de Gasto')
lista(doc, [
    'Estructura jerárquica (títulos y rubros hoja).',
    'Clasificación por sección y fuente.',
    'Campo de apropiación definitiva editable.',
])

h3(doc, '3.2.3 Importación Anexo 2 y Ejecución')
lista(doc, [
    'Importación del archivo "ANEXO 2 PRESUPUESTO DE GASTOS CENTRAL.xlsx".',
    'Importación de ejecución mensual ("EJECUCION GASTOS FEBRERO.xlsx" u otro mes).',
    'Edición manual de compromisos y pagos.',
])

h3(doc, '3.2.4 Deuda Pública')
lista(doc, [
    'Contratos de deuda con condiciones (entidad, valor, plazo, tasa).',
    'Pagarés asociados a cada contrato.',
    'Tabla de amortización automática por pagaré.',
    'Resumen consolidado del servicio a la deuda.',
])

h3(doc, '3.2.5 Costo de Personal')
lista(doc, [
    'Planta de cargos vigente con salarios y factores prestacionales.',
    'Cálculo automático del costo anual por cargo.',
    'Base para proyectar servicios personales asociados y contribuciones.',
])

h3(doc, '3.2.6 Vigencias Futuras')
lista(doc, [
    'Registro de vigencias futuras aprobadas o en ejecución.',
    'Clasificación por año, fuente y objeto.',
])

h3(doc, '3.2.7 Techos de Inversión')
parrafo(doc,
    'Reporte que muestra los topes de inversión según Ley 617/2000 y la asignación '
    'mínima obligatoria: 10% Predial destinado a Vivienda, destinaciones SGP (Educación, '
    'Salud, Agua Potable), entre otros.'
)

h3(doc, '3.2.8 Reportes y Exportación')
lista(doc, [
    'Reporte de gastos agrupado por sección y rubro.',
    'Reporte de ejecución mensual.',
    'Exportación a Excel.',
])

# -------- 3.3 Core
h2(doc, '3.3 Módulo Core')

h3(doc, '3.3.1 Dashboard')
parrafo(doc,
    'Pantalla principal con indicadores: total de ingresos proyectados, total de gastos, '
    'superávit/déficit y cumplimiento de indicadores Ley 617.'
)

h3(doc, '3.3.2 Tabla Concejo y Personería (Ley 617/2000)')
lista(doc, [
    'Límites de gasto de Concejo y Personería por categoría municipal.',
    'Número de concejales, sesiones ordinarias y extraordinarias, honorarios en factor SMLMV.',
    'Cálculo del honorario total del Concejo = honorario × sesiones × concejales.',
    'Cálculo del límite de Personería como % de ICLD.',
])

h3(doc, '3.3.3 Autenticación')
lista(doc, [
    'Login y logout de usuarios.',
    'Registro de nuevos usuarios con validación de contraseña.',
])

doc.add_page_break()

# ===== 4. FLUJO DE TRABAJO =====
h1(doc, '4. Flujo de Trabajo Recomendado')

parrafo(doc, 'Para elaborar el presupuesto de una vigencia, siga el orden:', negrita=True)

pasos = [
    ('1.', 'Configurar Parámetros del Sistema (UVT, IPC, PIB, % eficiencia, % crecimiento viviendas, categoría municipio, SMLMV).'),
    ('2.', 'Cargar Tabla Concejo/Personería si es primera vez.'),
    ('3.', 'Importar tarifas de Predial e ICA para la vigencia.'),
    ('4.', 'Importar Tabla Predial Comparativo (archivo XLSX del catastro).'),
    ('5.', 'Registrar contribuyentes ICA (manual o importación).'),
    ('6.', 'Cargar cartera de vigencias anteriores (por año).'),
    ('7.', 'Importar cifras históricas CUIPO 2022-2025 y calcular TCPA.'),
    ('8.', 'Importar Anexo 1 (estructura de rubros de ingreso).'),
    ('9.', 'Configurar el método de cálculo de cada rubro (PUVA, ICAI, IPC, etc.).'),
    ('10.', 'Ejecutar "Calcular Todos" — actualiza predial, ICA y apropiaciones.'),
    ('11.', 'Importar Anexo 2 (gastos), editar apropiaciones y registrar deuda.'),
    ('12.', 'Cargar planta de personal y vigencias futuras.'),
    ('13.', 'Revisar techos de inversión y límites Ley 617.'),
    ('14.', 'Generar y exportar los reportes de ingresos y gastos a Excel.'),
]
for n, txt in pasos:
    p = doc.add_paragraph()
    r = p.add_run(n + ' ')
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(txt).font.size = Pt(11)

# ===== 5. FÓRMULAS CLAVE =====
h1(doc, '5. Fórmulas Clave')

h2(doc, '5.1 Predial Vigencia Actual')
parrafo(doc,
    'Proyección = (Σ Avalúo × Tarifa / 1000) × % Eficiencia × (1 + % Crecimiento Viviendas)'
)
parrafo(doc,
    'El factor (1 + % Crecimiento Viviendas) aplica ÚNICAMENTE a la categoría UV '
    '(Urbano Vivienda). Las demás categorías usan factor 1.'
)

h2(doc, '5.2 Predial Vigencias Anteriores (Cartera)')
parrafo(doc,
    'Proyección = Valor Cartera × (% Base / 100) × (% Urbano o % Rural / 100)'
)

h2(doc, '5.3 ICA')
parrafo(doc, 'Ingresos Proyectados = Ingresos Brutos × (1 + PIB Nominal)')
parrafo(doc, 'Impuesto ICA = Ingresos Proyectados × Tarifa / 1000')
parrafo(doc, 'Avisos y Tableros = 15% × Total ICA')

h2(doc, '5.4 Métodos IPC / ICN')
parrafo(doc, 'Proyección = Recaudo Vigencia Anterior × (1 + Tasa IPC o ICN)')

h2(doc, '5.5 Método POAI')
parrafo(doc, 'Proyección = POAI Total Inversión × Tarifa POAI del rubro')

h2(doc, '5.6 Ley 617/2000 – Concejo y Personería')
parrafo(doc, 'Honorario Concejo = Honorario×SMLMV × (Sesiones Ord. + Extra) × N° Concejales')
parrafo(doc, 'Límite Concejo = ICLD × % Límite Concejo / 100')
parrafo(doc, 'Límite Personería = ICLD × % Límite Personería / 100')

# ===== 6. IMPORTACIONES / ARCHIVOS =====
h1(doc, '6. Archivos de Importación')

tabla_func(doc, [
    ('ANEXO 1 PRESUPUESTO DE INGRESOS 2026.xlsx', 'Estructura jerárquica de rubros de ingreso, hojas de cálculo predial y ICA de referencia.'),
    ('ANEXO 2 PRESUPUESTO DE GASTOS CENTRAL 2026.xlsx', 'Estructura de rubros de gasto de la administración central.'),
    ('Tabla Predial Comparativo con vigencia anterior.xlsx', 'Base catastral del municipio con avalúos vigencia actual vs anterior (13.729 predios aprox.).'),
    ('EJECUCION GASTOS FEBRERO 2026.xlsx', 'Ejecución presupuestal mensual de gastos.'),
    ('COSTO PLANTA VIGENTE 2026.xlsx', 'Planta de personal con salarios y factores.'),
])

parrafo(doc,
    'Los archivos deben colocarse en la raíz del proyecto (o en la carpeta "nuevos archivos"). '
    'El sistema detecta la estructura por encabezados y columnas conocidas.'
)

# ===== 7. MANTENIMIENTO =====
h1(doc, '7. Mantenimiento')

h2(doc, '7.1 Copias de seguridad')
lista(doc, [
    'El archivo db.sqlite3 contiene toda la información. Respaldar antes de importaciones masivas.',
    'Exportar los reportes a Excel al cierre de cada vigencia.',
])

h2(doc, '7.2 Actualización de vigencia')
lista(doc, [
    'Crear un nuevo registro en Parámetros del Sistema para la nueva vigencia y marcarlo activo.',
    'Importar las tarifas de Predial e ICA de la nueva vigencia.',
    'Actualizar valor UVT y SMLMV.',
    'Cargar cifras históricas del año que acaba de cerrar.',
])

h2(doc, '7.3 Migraciones de base de datos')
parrafo(doc, 'Después de actualizar el código:')
p = doc.add_paragraph()
r = p.add_run('python manage.py migrate')
r.font.name = 'Consolas'
r.font.size = Pt(11)

# ===== 8. GLOSARIO =====
h1(doc, '8. Glosario')
tabla_func(doc, [
    ('CUIPO', 'Clasificador Único de Ingresos y Gastos del orden territorial (CGN).'),
    ('ICLD', 'Ingresos Corrientes de Libre Destinación.'),
    ('SGP', 'Sistema General de Participaciones (transferencias de la Nación).'),
    ('POAI', 'Plan Operativo Anual de Inversiones.'),
    ('UVT', 'Unidad de Valor Tributario fijada por la DIAN.'),
    ('IPC', 'Índice de Precios al Consumidor.'),
    ('ICN', 'Ingresos Corrientes de la Nación.'),
    ('ICA', 'Impuesto de Industria y Comercio.'),
    ('SMLMV', 'Salario Mínimo Legal Mensual Vigente.'),
    ('TCPA', 'Tasa de Crecimiento Promedio Anual (compuesta).'),
    ('Avalúo catastral', 'Valor del predio determinado por el IGAC / catastro municipal.'),
    ('Tarifa por mil (‰)', 'Tarifa expresada en milésimas (‰) del avalúo o del ingreso base.'),
    ('Ley 617/2000', 'Ley colombiana que fija límites al gasto de funcionamiento territorial.'),
    ('Ley 44/1990', 'Ley del impuesto predial unificado.'),
    ('Ley 14/1983', 'Ley que regula el impuesto de Industria y Comercio.'),
    ('Vigencia', 'Año fiscal para el cual se presupuesta.'),
    ('Apropiación', 'Monto autorizado de gasto en el presupuesto.'),
    ('Rubro', 'Concepto presupuestal (ingreso o gasto).'),
])

# Pie con info del municipio
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\n\n\nAlcaldía de Puerto López — Departamento del Meta')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = AZUL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Secretaría de Hacienda')
r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Acuerdo No. 025 de 12-04-2020')
r.italic = True
r.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('República de Colombia')
r.font.size = Pt(10)

# Guardar
ruta = 'Manual_Usuario_Presupuesto.docx'
doc.save(ruta)
print(f'Manual generado: {ruta}')
